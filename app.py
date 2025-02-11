import streamlit as st 
from langchain_groq import ChatGroq
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from dotenv import load_dotenv
load_dotenv()
import streamlit as st
import googleapiclient.discovery
from gtts import gTTS
import tempfile
from playsound import playsound
from elevenlabs import ElevenLabs
import elevenlabs
import os,docx
# import pygame
import streamlit as st
from langchain_community.document_loaders import (
    TextLoader,
    PyPDFLoader,
    UnstructuredWordDocumentLoader
)
# from qdrant_client import QdrantClient
import os
import tempfile
import docx
import random
import string
from dotenv import load_dotenv
from datetime import datetime
# from streamlit_chat import message
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document
from langchain_community.vectorstores import Qdrant
from langchain_huggingface import HuggingFaceEmbeddings
from langchain.chains import create_retrieval_chain
from langchain_groq import ChatGroq
from langchain_core.prompts import ChatPromptTemplate

# Initialize the HuggingFace embeddings model
embedding_model = HuggingFaceEmbeddings(model_name="sentence-transformers/all-mpnet-base-v2")
# GROQ_API_KEY = st.secrets['GROQ_API_KEY']
# GROQ_API_KEY = "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"
# QDRANT_URL = st.secrets['QDRANT_URL']
# QDRANT_KEY =st.secrets['QDRANT_KEY'] 
# YOUTUB_API_KEY = st.secrets['YOUTUB_API_KEY']
# ELEVENLAB_API_KEY = st.secrets['ELEVENLAB_API_KEY']
# TAVILY_API_KEY =st.secrets['TAVILY_API_KEY'] 
# GOOGLE_API_KEY =st.secrets['GOOGLE_API_KEY'] 
GROQ_API_KEY = "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"
QDRANT_URL = "https://f757e150-6e65-48cd-a45a-ba46fa947234.eu-west-2-0.aws.cloud.qdrant.io:6333"
QDRANT_KEY = "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJhY2Nlc3MiOiJtIiwiZXhwIjoxNzQ2OTQwODUzfQ.Rp_WmX3FMJA3gJPgIi0g8Vp8Vn0MMPmoGGfeKaU8m0w"

# LANGSMITH_TRACING=True
LANGSMITH_ENDPOINT="https://api.smith.langchain.com"
LANGSMITH_API_KEY="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"
LANGSMITH_PROJECT="pr-new-mesenchyme-9"
# OPENAI_API_KEY="<your-openai-api-key>"
ELEVENLAB_API_KEY="sk_eaf6096f9230681058c5890f1de2d1b1577f2c8e94852a93"
TAVILY_API_KEY="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"
OPENAI_API_KEY="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"
YOUTUB_API_KEY="AIzaSyChHqpzh0m34PHsRYfRp4dGC92YPFFgbks"  # Replace with your actual API key"
GOOGLE_API_KEY="AIzaSyBjnj_9B9CLmaJGueSVWJHGwz7QEC6TW9M"
# Initialize the LLM
chat_llm = ChatGroq(model_name="llama-3.3-70b-versatile", api_key=GROQ_API_KEY)
# Set your Groq API key from environment variable

if not GROQ_API_KEY:
    raise ValueError("GROQ_API_KEY not set in environment variables")
# Initialize chat history in session state
if "chat_history_rag" not in st.session_state:
    st.session_state.chat_history_rag = []

# Define the prompt template
prompt_str = """Use the following retrieved context to answer the question:
Context: {context}
Conversation History: {chat_history_rag}
Question: {question}"""
prompt = ChatPromptTemplate.from_template(prompt_str)



# Function to process uploaded files
def process_uploaded_files(uploaded_files):
    text_chunks_list = []
    for uploaded_file in uploaded_files:
        file_name = uploaded_file.name
        file_text = extract_text_from_file(uploaded_file)
        text_chunks = split_text_into_chunks(file_text, file_name)
        text_chunks_list.extend(text_chunks)
    return text_chunks_list

# Function to extract text from files
def extract_text_from_file(uploaded_file):
    text = ""
    file_extension = os.path.splitext(uploaded_file.name)[1].lower()
    with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as temp_file:
        temp_file.write(uploaded_file.getvalue())
        temp_path = temp_file.name
    if file_extension == ".pdf":
        loader = PyPDFLoader(temp_path)
        pages = loader.load()
        text = "\n".join([page.page_content for page in pages])
    elif file_extension == ".docx":
        doc = docx.Document(temp_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    elif file_extension == ".txt":
        with open(temp_path, "r", encoding="utf-8") as file:
            text = file.read()
    os.unlink(temp_path)
    return text

# Function to split text into chunks
def split_text_into_chunks(file_text, file_name):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=780, chunk_overlap=80, length_function=len)
    text_chunks = text_splitter.split_text(file_text)
    doc_list = []
    for text_chunk in text_chunks:
        file_source = {"source": file_name}
        doc_in_str = Document(page_content=text_chunk, metadata=file_source)
        doc_list.append(doc_in_str)
    return doc_list

# Function to create a vector store
def create_vector_store(text_chunks):
    vectorstore = Qdrant.from_documents(
        documents=text_chunks,
        embedding=embedding_model,
        url=QDRANT_URL,
        api_key=QDRANT_KEY
    )
    return vectorstore

# Function to handle user input
def handle_user_input(user_question, vectorstore):
    # Retrieve relevant context from Qdrant
    retrieved_chunks = vectorstore.similarity_search(user_question, k=3)
    # st.write(retrieved_chunks)
    source=retrieved_chunks[0].metadata['source']
    context_text = "\n\n".join([chunk.page_content for chunk in retrieved_chunks])
    # st.write(context_text)
   
    # Prepare input data for the chain
    input_data = {
        "question": user_question,
        "context": context_text,
        "chat_history_rag": st.session_state.chat_history_rag,
    }
    # Generate response
    response = prompt | chat_llm
    ai_response = response.invoke(input_data)
    ai_response= ai_response.content
    # Update chat history
    st.session_state.chat_history_rag.append({"role":"user","content":user_question})
    st.session_state.chat_history_rag.append({"role":"assistant","content":f"{ai_response} \n\n     Source: {source}" })
    # Display chat history
    for  message in st.session_state.chat_history_rag:
         st.chat_message(message['role']).markdown(message['content'])
# Streamlit app layout
def RAG():
    # st.set_page_config("Document QA System", page_icon="📚")
    st.header("📚 Retrieve | Augment | Generate")
    with st.sidebar:
        st.title("📚 Documents QA System")
        st.markdown("""
        Upload a Documents and ask questions about its content.
        **Features:**
        - Documents text extraction
        - AI-powered answers
        - Source tracking
        """)
        uploaded_files = st.file_uploader("Upload Here:", type=['pdf', 'docx', 'txt'], accept_multiple_files=True)
        process = st.button("Start Process")
    if process:
        if not uploaded_files:
            st.warning("Please upload files first!")
            st.stop()
        if not GROQ_API_KEY:
            st.warning("Please provide your API key first!")
            st.stop()
        text_chunks_list = process_uploaded_files(uploaded_files)
        vectorstore = create_vector_store(text_chunks_list)
        st.write("Vector Store Created.")
        st.session_state.vectorstore = vectorstore
    if "vectorstore" in st.session_state:
        user_question = st.chat_input("Ask questions about your files:")
        if user_question:
            handle_user_input(user_question, st.session_state.vectorstore)



def search_youtube(query, max_results=2):
  
    youtube = googleapiclient.discovery.build("youtube", "v3", developerKey=YOUTUB_API_KEY)

    try:
        request = youtube.search().list(
            part="snippet",
            q=query + " in channel muhammadirfanmalik",
            type="video",
            maxResults=max_results,
            order="relevance",  # Can be: 'date', 'rating', 'viewCount', 'relevance'
            fields="items(id(videoId))"  # Only get required fields to save quota
        )
    except googleapiclient.errors.HttpError as e:
        st.error(f"An API error occurred: {str(e)}")
        return []

    response = request.execute()

    video_urls = []
    if response.get('items'):
        for item in response['items']:
            video_id = item['id']['videoId']
            video_urls.append(f"https://www.youtube.com/watch?v={video_id}")
        return video_urls
    else:
        return ["No videos found for this query in the specified channel."]

def speak_text_using_gtts(response):
    try:
        tts=gTTS(response)
        with tempfile.NamedTemporaryFile(delete=False,suffix=".mp3") as fp:
            temporary_path=fp.name
            tts.save(temporary_path)
        playsound(temporary_path)
        os.remove(temporary_path)
    except Exception as e:
        st.write("An Error occured while TTS",e)
    
def speak_text_using_elevanlab(response):
        # Initialize the ElevenLabs client using the API key from environment variables.
    client = ElevenLabs(api_key=ELEVENLAB_API_KEY)
    
    # Generate audio from the provided response text using the specified parameters.
    audio = client.generate(
        text=response,
        voice="Antoni",                   # Specify the voice (e.g., "Aria")
        output_format="mp3_44100_128",    # Use high-quality MP3 format
        model="eleven_turbo_v2"           # Specify the model version
    )
    try:
        #save the generated audio to a temporary file.
        with tempfile.NamedTemporaryFile(delete=False,suffix=".mp3") as fp:
            temporary_path=fp.name 
            #use elevenlabs save function to write the audio to our temporaray path
            elevenlabs.save(audio,temporary_path) 
        playsound(temporary_path)
        os.remove(temporary_path)
    except Exception as e:
        st.error("An Error Occured",e)
        
def irfan_malik():
    system_prompt="""
            Your name is Irfan Malik. You are an accomplished Artificial Intelligence and Data Science expert with millions in sales, a top freelancer, and the founder and CEO of Xeven Solutions. Your company, Xeven Solutions, is a renowned AI and software development firm that serves over 500 clients worldwide and employs more than 200 professionals. 

            At Xeven Solutions, we specialize in delivering cutting-edge technology solutions, including:
            1. AI Development Services – Encompassing Machine Learning, Deep Learning, Fine-tuning, RAG, Natural Language Processing, Computer Vision, and Predictive Modeling.
            2. Mobile App Development – Crafting innovative, user-friendly mobile applications.
            3. Custom Web Development – Developing bespoke web solutions tailored to clients' unique business needs.
            4. Data Analytics – Providing actionable insights through advanced data processing and visualization.

            Our company is recognized as a trusted AI development partner, building meaningful AI healthcare solutions, intelligent chatbots, and offering services such as ChatGPT integration, custom software development, digital marketing, DevOps, UI/UX design, and much more.

            As Irfan Malik, you are here to educate and guide people by answering their queries with wisdom and clarity. Always respond in a polite, professional manner that reflects the mindset of a millionaire CEO who values excellence, innovation, and customer success.
            Make the response concise and clear.
            
            minimum lines sould be 3 and maximum 10
            Your motive is focus,attention to detail,consistancy and then hard work.
            
            Question: {question}
            # As I'm also getting a video from his youtube chennel so always say 'here is the video about' in the end in a new line with user's actual question in [].
            """

    st.title("Irfan Malik CEO Of Xeven Solutions")

    if "chat_history_irfan" not in st.session_state:
        st.session_state.chat_history_irfan=[]
        
        
    prompt =ChatPromptTemplate.from_template(system_prompt)
    

    query=st.chat_input("Ask Anything.")
    if query:
        # st.chat_message("user").markdown(query)
        st.session_state.chat_history_irfan.append({"role":"user","type":"text","content":query}) 
        chain= prompt | chat_llm | StrOutputParser()

        response=chain.invoke({"question":query})
        # st.chat_message('assistant').markdown(response)
        st.session_state.chat_history_irfan.append({"role":"assistant","type":"text","content":response})
    
    
        #getting youtube url
        video_urls=search_youtube(query)
        if video_urls and video_urls[0] != "No videos found for this query in the specified channel.":
            for url in video_urls:
                # with st.chat_message(assi)
                # st.video(url)
                st.session_state.chat_history_irfan.append({
                    "role":"user",
                    "type":"video",
                    "content":url
                })
        else:
            st.write(video_urls[0])
            

    for message in st.session_state.chat_history_irfan:
        if message['role']=='user' and message["type"]=="text":
            with st.chat_message('human'):
                # if message["type"]=="text":
                    st.markdown(message["content"])
        elif message['role']=='user' and  message["type"]=="video":
            with st.chat_message("assistant"):
                st.video(message["content"])
        elif message["role"] =="assistant":
            with st.chat_message('ai'):
                st.markdown(message["content"])
                # speak_text_using_gtts(message['content'][-1])
                    

    # if st.session_state.chat_history_irfan:
    #     last_message=st.session_state.chat_history_irfan[-3]
    #     if last_message['role']=='assistant':
    #         #speak
    #         # st.write(last_message["content"])
    #         # pass 
    #         # speak_text_using_elevanlab(last_message['content'])
    #         speak_text_using_gtts(last_message["content"])
    #     else:
    #         st.error("last messsasge is not ai")
    # else:
    #     st.write("no chat history")
    
def my_skills():
    st.title("Muhammad Ahtisham Yousaf")
    st.subheader("Asslam-u-Alaikum!Sir this project is created just for showing my skills on Sunday and Monday.")
    

from groq import Groq
import streamlit as st 
import os
import base64
from gtts import gTTS
# import vlc
import time
from dotenv import load_dotenv
load_dotenv()


vision_llm = Groq(api_key=GROQ_API_KEY)
      
from gtts import gTTS
from playsound import playsound
import tempfile
import pygame

# Function to Convert Text to Speech
# Purpose: Initializes the pygame audio system.
# Why: Without this, pygame can’t play audio.
# Note: This should be called before using any pygame audio functions.
# pygame.mixer.init()

# Function to Convert Text to Speech and Save Temporarily
# def generate_audio(text):
#     try:
#         tts = gTTS(text)
#         with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as fp:
#             temp_path = fp.name
#             tts.save(temp_path)
#         return temp_path  # Return the path of the audio file
#     except Exception as e:
#         print(f"Error in TTS: {e}")
#         return None

# # Function to Play Audio
# def play_audio(audio_path):
#     try:
#         pygame.mixer.music.load(audio_path)
#         pygame.mixer.music.play()
#     except Exception as e:
#         print(f"Error playing audio: {e}")

# # Function to Stop Audio
# def stop_audio():
#     pygame.mixer.music.stop()
            
# STAGE 2: Convert image to required format
def encode_image(uploaded_image):
    # try:
    #     with open(uploaded_image, "rb") as image_file:
    #         encoded_image = base64.b64encode(image_file.read()).decode('utf-8')
    #         return encoded_image
    # except Exception as e:
    #     raise RuntimeError(f"Failed to read or encode image: {str(e)}")
    try:
        # Read bytes directly from the uploaded file
        encoded_image = base64.b64encode(uploaded_image.read()).decode('utf-8')
        return encoded_image
    except Exception as e:
        raise RuntimeError(f"Failed to read or encode image: {str(e)}")
# Correct data URL prefix (note the semicolon after 'jpeg')
# data_url = f"data:image/jpeg;base64,{encoded_image}"
# encoded_image=encode_image(image_path)
# STAGE 3: Setting up the multimodal (vision) LLM

# query = "what is the problem with my face"
# model = "llama-3.2-90b-vision-preview"

def analyze_image_with_query(query,model,encoded_image):
# Construct messages list with both text and image content
    messages = [
        {
            "role": "user",
            "content": [
                {"type": "text", "text": query},
                {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{encoded_image}"}}
            ]
        }
    ]

    # Generate chat completion using the Groq API
    try:
        chat_completion = vision_llm.chat.completions.create(messages=messages, model=model)
        return chat_completion.choices[0].message.content
    except Exception as e:
        return f"Error generating chat completion: {e}"

# Main function
def vision():
    st.title("🖼️ Computer Vision Assistant")

    # Session state for chat history
    if "chat_history_vision" not in st.session_state:
        st.session_state.chat_history_vision = []
    
    if "process_complete_vision" not in st.session_state:
        st.session_state.process_complete_vision=False
    
    if "audio_paths" not in st.session_state:
        st.session_state.audio_paths = {}  # To map responses with their audio files

    with st.sidebar:
    # Upload Image
        uploaded_image = st.file_uploader("📤 Upload an image here:")
        if uploaded_image:
            # Encode image for API
            encoded_image = encode_image(uploaded_image)
            st.session_state.process_complete_vision = True

    if st.session_state.process_complete_vision:
        # User Prompt
        query = st.chat_input("💬 Your Prompt:")
        if query:
            # Append user query
            st.session_state.chat_history_vision.append({
                "role": "user",
                "type": "text",
                "content": query
            })

            # Process with Groq API
            model = "llama-3.2-90b-vision-preview"
            response = analyze_image_with_query(query, model, encoded_image)
            # Append image to chat history
            st.session_state.chat_history_vision.append({
                "role": "user",
                "type": "image",
                "content": uploaded_image
            })
            # Append assistant response
            st.session_state.chat_history_vision.append({
                "role": "assistant",
                "type": "text",
                "content": response
            })

            #  # Generate and save audio for the response
            # audio_path = generate_audio(response)
            # if audio_path:
            #     st.session_state.audio_paths[response] = audio_path
    # Render Chat History (Simplified)
    for message in st.session_state.chat_history_vision:
        if message['role']=="user":
            with st.chat_message('user'):
                if message["type"]== "text":
                    st.markdown(message["content"])
                elif message['type']=="image":
                    st.image(message['content'])
        elif message["role"]=="assistant":
            with st.chat_message("assistant"):
                st.markdown(message["content"])
                # speak_text(message['content'])

                # # Speak/Stop Button
                # is_playing = st.session_state.get(f"is_playing_{message['content']}", False)
                # if st.button("🎵 Read Load" if not is_playing else "⏹️ Stop", key=f"play_{message['content']}"):
                #     if not is_playing:
                #         play_audio(st.session_state.audio_paths[message['content']])
                #         st.session_state[f"is_playing_{message['content']}"] = True
                #     else:
                #         stop_audio()
                #         st.session_state[f"is_playing_{message['content']}"] = False

                # 🎵 Audio Download Button
                # audio_path = st.session_state.audio_paths.get(message["content"])
                # if audio_path:
                #     with open(audio_path, "rb") as audio_file:
                #         audio_data = audio_file.read()
                #         st.download_button(
                #             label="⬇️ Download Audio",
                #             data=audio_data,
                #             file_name="response_audio.mp3",
                #             mime="audio/mpeg",
                #             key=f"download_{message['content']}"
                #         )
                        


# Function to reset session state when option changes
def reset_session():
    if "last_selection" in st.session_state and st.session_state.last_selection != selection:
        st.session_state.clear()
        st.session_state.last_selection = selection
        st.rerun()  # Rerun Streamlit to clear everything
        st.write("Reruning the app")

# Sidebar selection
with st.sidebar:
    selection = st.radio("Select Your Preference", options=["My Skills", 'RAG', 'Sir Irfan Malik', "Vision", "AI Agent"])

# Reset session state if selection changes
if "last_selection" not in st.session_state:
    st.session_state.last_selection = selection
reset_session()


if selection=="Sir Irfan Malik":
    irfan_malik()
elif selection=='My Skills':
    my_skills()
elif selection=='RAG':
    RAG()    
elif selection=='Vision':
    vision()
elif selection=="AI Agent":
    from dotenv import load_dotenv
    from langchain_groq import ChatGroq
    from langchain_community.chat_models import ChatOpenAI
    from langchain_community.tools import TavilySearchResults
    from langgraph.prebuilt import create_react_agent
    from langchain_core.messages.ai import AIMessage
    load_dotenv()
    import os 

    # Phase1–Create AI Agent
    # # 1. Setup API Keys for Groq and Tavily
    # GROQ_API_KEY=os.environ.get("GROQ_API_KEY")
    # TAVILY_API_KEY=os.environ.get("TAVILY_API_KEY")
    # OPENAI_API_KEY=os.environ.get("OPENAI_API_KEY")

    # 2. Setup LLM & Tools

    # from tavily import TavilyClient
    # tavily_client = TavilyClient(api_key=TAVILY_API_KEY)
    # response = tavily_client.search("Who is Leo Messi?")
    # print(response)

    groq_llm=ChatGroq(model_name="llama-3.3-70b-versatile")
    # openai_llm=ChatOpenAI(model="gpt-4o-mini")


    # 3. Setup AI Agent with Search tool functionality

    system_prompt="Act as an chatbot who is smart and friendly."
    from langchain_google_genai import ChatGoogleGenerativeAI as genai
    def get_response_from_ai_agent(my_model,query,allow_search,system_prompt,provider):
        if provider=="Groq":
            llm=ChatGroq(model_name=my_model)
        elif provider =="Google":
            llm=genai(model=my_model,api_key=GOOGLE_API_KEY)
        
        
        
        search_tools=[TavilySearchResults(max_results=2)] if allow_search else []
        agent = create_react_agent(
            model=llm,  # Ensure this is valid
            tools=search_tools,
            state_modifier=system_prompt  # Correct argument name
        )

        # query="Tell me about the trends in Agentic AI" 
        state={"messages":query}
        response=agent.invoke(state)
        messages=response.get("messages")
        ai_response=[message.content for message in messages if isinstance(message,AIMessage)]
        return ai_response[-1]

    # Phase2–Setup Backend (With FastAPI)
    from pydantic import BaseModel
    from fastapi import FastAPI


    # 1. Setup Pydantic Model (Schema Validation)

    class RequestState(BaseModel):
        model_name:str
        model_platform:str 
        system_prompt:str 
        messages:list[str]
        allow_search:bool 
        
        
        class Config:
            protected_namespaces = ()

    # 2. Setup AI Agent from FrontEnd Request

    ALLOWED_MODEL_NAMES=["llama3-70b-8192", "mixtral-8x7b-32768", "llama-3.3-70b-versatile", "gemini-1.5-pro"]
    app=FastAPI(title="LangGraph AI AGent")

    @app.post("/chat")
    def chat_endpoint(request: RequestState):
        """
        API Endpoint to interact with the chatbot using langgraph and search tools.
        it dynamically selects the model specified in the request
        """
        if request.model_name not in ALLOWED_MODEL_NAMES:
            return {"error":"Invalid model name.Kindly select a valid AI Model"}
        
        my_model=request.model_name
        query=request.messages
        allow_search=request.allow_search
        system_prompt=request.system_prompt
        platform=request.model_platform
        
        #Create AI Agent and get response from it
        response=get_response_from_ai_agent(my_model,query,allow_search,system_prompt,platform)
        # print(response)
        return response
        
    # Phase3–Setup Frontend
    # 1. Setup UI with streamlit (model provider, model, system prompt, query)
    import streamlit as st 

    st.title("AI ChatBot Agent")
    st.write("Create and Interact with the AI Agents")

    if "chat_history_agent" not in st.session_state:
        st.session_state.chat_history_agent=[]

    with st.sidebar:
        
        system_prompt=st.text_area("Define Role for YOur AI AGent",height=70,placeholder="Write here your Agent role")

        model_platform=st.radio("Select Platform:",("Groq","Google"))

        GROQ_MODEL_NAMES=["llama-3.3-70b-versatile","mixtral-8x7b-32768"]
        GOOGLE_MODEL_NAMES=["gemini-1.5-pro"]

        if model_platform=="Groq":
            selected_model=st.selectbox("Select Groq Model:",GROQ_MODEL_NAMES)
        elif model_platform=="Google":
            selected_model=st.selectbox("Select Google AI Model:",GOOGLE_MODEL_NAMES)

        allow_web_search=st.checkbox("Allow Web Search")

    user_query=st.chat_input("Ask Anything!")

    FASTAPI_URL="http://127.0.0.1:9999/chat"


    if user_query:
        user_query=user_query.strip()
        # st.chat_message('human').markdown(user_query)
        st.session_state.chat_history_agent.append({"role":"user","content":user_query})
        #Step 2: Connect with backend via url 
        import requests
        
        payload={
            "model_name":selected_model,
            "model_platform":model_platform,
            "system_prompt":system_prompt,
            "messages":[user_query],
            "allow_search": allow_web_search
        }
        response=requests.post(FASTAPI_URL,json=payload)
        if response.status_code==200:
            response_data=response.json()
            st.session_state.chat_history_agent.append({"role":"assistant","content":response_data})
            if "error" in response_data:
                st.write(response_data["error"])
        
    for message in st.session_state.chat_history_agent:
        st.chat_message(message["role"]).markdown(message["content"])




    # 3. Run app & Explore Swagger UI Docs
    if __name__=="__main__":
        import uvicorn
        uvicorn.run(app,host="127.0.0.1",port=9999)

   

    
    
